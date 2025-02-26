import hashlib
import subprocess
from docx import Document
import os
import re
import requests
from concurrent.futures import ThreadPoolExecutor
import logging
import time
import json
from urllib.parse import quote

# 配置日志记录
logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s - %(levelname)s - %(message)s')

CONFIG = {
    "translate_api": "baidu",  # microsoft/baidu/deepl
    "baidu_appid": "",
    "baidu_secret": "",
    "proxy": None,  # 例如："http://127.0.0.1:1080"
    "max_workers": 2,
    "request_timeout": 15,
    "baidu_max_length": 6000,
    "request_timeout": (15, 30),
    "retry_delay": 1.0,
    "max_text_length": 4000
}


def split_text(text):
    """改进的分割函数"""
    # 添加编码长度校验
    encoded_length = len(quote(text))
    if encoded_length > CONFIG["baidu_max_length"]:
        # 按实际编码长度分割
        return split_by_encoded_length(text)
    return [text]


def split_by_encoded_length(text):
    """按编码后长度分割"""
    chunks = []
    current_chunk = []
    current_length = 0

    for char in text:
        encoded_char = quote(char)
        char_length = len(encoded_char)
        if current_length + char_length > CONFIG["baidu_max_length"]:
            chunks.append(''.join(current_chunk))
            current_chunk = [char]
            current_length = char_length
        else:
            current_chunk.append(char)
            current_length += char_length

    if current_chunk:
        chunks.append(''.join(current_chunk))
    return chunks


def translate_text(text, src_lang='en', target_lang='zh-CN', max_retries=3):
    """翻译函数"""
    if not text.strip():
        return ""

    for attempt in range(max_retries):
        try:
            if CONFIG["translate_api"] == "baidu":
                return translate_baidu(text, src_lang, target_lang)
            else:
                raise ValueError("不支持的翻译API")

        except Exception as e:
            logging.warning(f"翻译尝试 {attempt + 1} 失败: {str(e)}")
            time.sleep(2 ** attempt)
    return "[翻译失败]"

LAST_REQUEST_TIME = 0
REQUEST_INTERVAL = 1.0

def translate_baidu(text, src_lang, target_lang):
    global LAST_REQUEST_TIME
    elapsed = time.time() - LAST_REQUEST_TIME
    if elapsed < REQUEST_INTERVAL:
        time.sleep(REQUEST_INTERVAL - elapsed)
    LAST_REQUEST_TIME = time.time()
    # 参数验证
    if target_lang.lower() == 'zh-cn':
        target_lang = 'zh'  # 百度使用zh表示简体中文

    valid_langs = ['zh', 'en', 'jp', 'kor', 'fra', 'ru']
    if target_lang not in valid_langs:
        raise ValueError(f"百度不支持的目标语言: {target_lang}")

    # 自动编码处理
    encoded_text = quote(text, safe='')  # 强制编码特殊字符
    if len(encoded_text) > 6000:
        raise ValueError("文本长度超过百度API限制")
    """使用百度翻译API"""
    url = "https://fanyi-api.baidu.com/api/trans/vip/translate"
    salt = str(time.time())
    sign_str = CONFIG["baidu_appid"] + text + salt + CONFIG["baidu_secret"]
    sign = hashlib.md5(sign_str.encode()).hexdigest()

    params = {
        "q": text,
        "from": src_lang,
        "to": target_lang,
        "appid": CONFIG["baidu_appid"],
        "salt": salt,
        "sign": sign
    }

    session = requests.Session()
    if CONFIG["proxy"]:
        session.proxies = {"https": CONFIG["proxy"]}

    try:
        response = session.get(  # 改为GET请求
            url,
            params=params,  # 使用查询参数
            timeout=(10, 30)
        )
        result = response.json()
        if "error_code" in result:
            error_code = result["error_code"]
            error_msg = result.get("error_msg", "未知错误")
            if error_code in ['54003', '54004', '54005']:  # 频率限制类错误
                time.sleep(2)
                return translate_baidu(text, src_lang, target_lang)
            elif error_code in ['52001', '52002', '52003']:  # 网络相关错误
                time.sleep(1)
                return translate_baidu(text, src_lang, target_lang)
            else:
                raise ValueError(f"百度API错误 {error_code}: {error_msg}")

        if "trans_result" not in result:
            raise ValueError("响应缺少trans_result字段")

        return "\n".join([item["dst"] for item in result["trans_result"]])
    except requests.exceptions.Timeout:
        logging.warning("百度API请求超时，尝试降低并发量")
        return translate_text(text)


def pdf_to_docx_ocr_translate(pdf_path, docx_path, ocr_lang="eng", target_lang="zh-CN",
                              dpi=300, poppler_path=None, tesseract_path=None):
    """
    PDF转Word文档函数
    """
    # OCR语言映射（扩展更多语言支持）
    lang_map = {
        "eng": "en",
        "chi_sim": "zh-CN",
        "jpn": "ja",
        "deu": "de",
        "fra": "fr",
        "spa": "es"
    }

    src_lang = lang_map.get(ocr_lang.lower(), ocr_lang.split('_')[0])

    # 配置路径
    env = os.environ.copy()
    if poppler_path:
        env["PATH"] = f"{poppler_path}{os.pathsep}{env['PATH']}"

    tesseract_cmd = tesseract_path if tesseract_path else "tesseract"
    # 创建临时目录
    temp_dir = "temp_ocr_images"
    os.makedirs(temp_dir, exist_ok=True)

    try:
        # 修改 PDF 转图片命令（添加 -rx 300 -ry 300 参数）
        poppler_command = [
            "pdftoppm",
            "-r", str(dpi),
            "-rx", "300",  # 添加水平分辨率
            "-ry", "300",  # 添加垂直分辨率
            "-png",
            pdf_path,
            os.path.join(temp_dir, "page")
        ]

        result = subprocess.run(
            poppler_command,
            capture_output=True,
            text=True,
            env=env,
            timeout=60
        )

        if result.returncode != 0:
            logging.error(f"PDF转图片失败: {result.stderr}")
            return

        # 获取生成的图片列表
        image_list = sorted([
            os.path.abspath(os.path.join(temp_dir, f))
            for f in os.listdir(temp_dir)
            if f.endswith(".png")
        ], key=lambda x: int(re.search(r'page-(\d+)', x).group(1)))

        if not image_list:
            logging.error("未生成任何图片文件")
            return

        logging.info(f"成功生成 {len(image_list)} 张图片")

        # 创建Word文档
        doc = Document()
        doc.add_heading("OCR翻译文档", 0)

        # 使用线程池加速OCR处理
        with ThreadPoolExecutor(max_workers=CONFIG["max_workers"]) as executor:
            futures = []
            for img_path in image_list:
                future = executor.submit(
                    process_image,
                    img_path,
                    tesseract_cmd,
                    ocr_lang,
                    src_lang,
                    target_lang
                )
                futures.append((img_path, future))

            for img_path, future in futures:
                try:
                    ocr_text, translated_text = future.result()
                    add_to_document(doc, img_path, ocr_text, translated_text)
                except Exception as e:
                    logging.error(f"处理 {img_path} 失败: {str(e)}")
                    doc.add_paragraph(f"处理失败: {img_path}")

        # 保存文档
        doc.save(docx_path)
        logging.info(f"文档已保存至: {docx_path}")

    finally:
        # 清理临时文件
        for img_path in image_list:
            try:
                os.remove(img_path)
            except Exception as e:
                logging.warning(f"删除临时文件失败 {img_path}: {str(e)}")
        try:
            os.rmdir(temp_dir)
        except Exception as e:
            logging.warning(f"删除临时目录失败: {str(e)}")

    # 在执行 OCR 前添加验证
    if not os.path.exists(img_path):
        raise FileNotFoundError(f"图片文件不存在: {img_path}")

    #    在调用 subprocess 前添加路径打印
    logging.info(f"正在处理图片路径: {img_path}")


def process_image(img_path, tesseract_cmd, ocr_lang, src_lang, target_lang):
    """处理单张图片的OCR和翻译"""
    try:
        # OCR识别
        cmd = [
            tesseract_cmd,
            img_path,
            "stdout",
            "-l", ocr_lang,
            "--psm", "6",
            "--oem", "3",
            "-c", "preserve_interword_spaces=1",
            "-c", "tessedit_char_whitelist=ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789.,;:!?()%-'\""
        ]

        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            encoding="utf-8",
            timeout=30
        )

        ocr_text = result.stdout.strip() if result.returncode == 0 else ""
        logging.info(f"OCR完成: {img_path} [长度: {len(ocr_text)}]")

        if not ocr_text:
            return "[OCR无内容]", "[无需翻译]"

        # 翻译文本
        translated_text = translate_text(ocr_text, src_lang, target_lang)
        return ocr_text, translated_text

    except Exception as e:
        raise RuntimeError(f"处理图片时出错: {str(e)}") from e


def add_to_document(doc, img_path, ocr_text, translated_text):
    """将内容添加到Word文档"""
    doc.add_heading(os.path.basename(img_path), level=2)

    # 添加原文段落
    doc.add_paragraph("原文内容：")
    para = doc.add_paragraph(ocr_text)
    para.style = "Body Text"

    # 添加翻译段落
    doc.add_paragraph("翻译内容：")
    para = doc.add_paragraph(translated_text)
    para.style = "Body Text"

    # 添加分页符
    doc.add_page_break()


# 添加密钥验证函数
def validate_baidu_credentials():
    test_params = {
        "q": "test",
        "from": "en",
        "to": "zh",
        "appid": CONFIG["baidu_appid"],
        "salt": "123456",
        "sign": hashlib.md5((CONFIG["baidu_appid"] + "test" + "123456" + CONFIG["baidu_secret"]).encode()).hexdigest()
    }

    response = requests.get("https://fanyi-api.baidu.com/api/trans/vip/translate", params=test_params)
    result = response.json()
    if "error_code" in result and result["error_code"] in ['52001', '54000']:
        raise ValueError("无效的API凭证，请检查appid和secret")

if __name__ == "__main__":
    try:
        validate_baidu_credentials()
    except Exception as e:
        logging.error(f"API凭证验证失败: {str(e)}")
        exit(1)
    # 配置路径（根据实际情况修改）
    input_pdf = r"./NED-Report-May-June-part-01.pdf"
    output_docx = r"./NED-Report-May-June-part-01.docx"
    poppler_path = r"D:/poppler-24.08.0/Library/bin"
    tesseract_path = r"D:/Tesseract-OCR/tesseract.exe"

    # 执行转换
    pdf_to_docx_ocr_translate(
        pdf_path=input_pdf,
        docx_path=output_docx,
        ocr_lang="eng",  # 中文文档使用 "chi_sim"
        target_lang="zh-CN",
        dpi=300,
        poppler_path=poppler_path,
        tesseract_path=tesseract_path
    )
